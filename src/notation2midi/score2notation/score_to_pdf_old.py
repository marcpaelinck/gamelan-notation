import os
import pickle
import time
import typing
from datetime import datetime
from enum import Enum
from os import path
from random import randint

import docx2pdf
from docx import Document
from docx.document import Document as DocxDocument
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import (
    WD_ALIGN_PARAGRAPH,
    WD_COLOR_INDEX,
    WD_TAB_ALIGNMENT,
    WD_TAB_LEADER,
)
from docx.oxml.parser import parse_xml
from docx.oxml.xmlchemy import BaseOxmlElement
from docx.section import Section
from docx.shared import Cm, Inches, Length, Pt, RGBColor
from docx.styles.style import BaseStyle, CharacterStyle, ParagraphStyle, _TableStyle
from docx.table import Table, _Cell, _Row
from docx.text.paragraph import Paragraph
from PIL import Image, ImageDraw, ImageFont

from src.common.classes import Gongan, Note, Score
from src.common.constants import DEFAULT, InstrumentType, Position, Stroke
from src.common.metadata_classes import DynamicsMeta, MetaData, MetaDataBaseModel
from src.notation2midi.classes import ParserModel
from src.notation2midi.score2notation.score_to_notation import aggregate_positions

# The following functions save a human readable version of the score to a PDF file.


class MergeType(Enum):
    RANGE = 1
    LAST_CELL = 2


class ScoreToPDFConverter(ParserModel):
    TAG_COLWIDTH = Cm(2.3)
    basicparastyle = None
    metadatastyle = None

    def __init__(self, score: Score, pickle: bool = False):
        super().__init__(self.ParserType.SCORETOPDF, score.settings)
        self.pdf_settings = self.run_settings.pdf_converter
        self.doc: DocxDocument = Document(os.path.join(self.pdf_settings.folder, self.pdf_settings.docx_template))
        self.score = score
        self.pickle = pickle
        self._format_notation()
        self.defaultparastyle = self.basicparastyle
        self.defaultcharstyle = self.metadatastyle
        fpath, ext = os.path.splitext(score.settings.notation.midi_out_filepath)
        self.filepath = fpath + ".docx"
        self.current_para = None
        self.current_tempo = -1
        self.current_dynamics = -1

    def _create_para_style(
        self,
        name: str,
        space_after: int = 0,
        keep_with_next: bool = False,
        hanging_indent: int = None,
        tab_stops: list[tuple[int, WD_TAB_ALIGNMENT, WD_TAB_LEADER]] = [],
        align: WD_ALIGN_PARAGRAPH = WD_ALIGN_PARAGRAPH.LEFT,
    ) -> ParagraphStyle:
        parastyle: ParagraphStyle = self.doc.styles.add_style(name, style_type=WD_STYLE_TYPE.PARAGRAPH, builtin=False)
        parastyle.base_style = self.doc.styles["Normal"]
        parastyle.font.name = "Arial"
        parastyle.paragraph_format.space_after = space_after
        parastyle.paragraph_format.keep_with_next = keep_with_next
        parastyle.paragraph_format.tab_stops.clear_all()
        if hanging_indent:
            parastyle.paragraph_format.left_indent = hanging_indent
            parastyle.paragraph_format.first_line_indent = -hanging_indent
        for tab, tabtype, leader in tab_stops:
            parastyle.paragraph_format.tab_stops.add_tab_stop(tab, alignment=tabtype, leader=leader)
        parastyle.font.highlight_color = WD_COLOR_INDEX.WHITE
        parastyle.paragraph_format.alignment = align

        return parastyle

    def _create_char_style(
        self,
        name: str,
        base: ParagraphStyle = None,
        font: str = "Arial",
        fsize: int = 10,
        bold=False,
        italic=False,
        rgb: tuple[int] = (0, 0, 0),
        hilite: WD_COLOR_INDEX = WD_COLOR_INDEX.WHITE,
    ) -> CharacterStyle:
        charstyle: CharacterStyle = self.doc.styles.add_style(name, style_type=WD_STYLE_TYPE.CHARACTER, builtin=False)
        charstyle.base_style = base or self.basicparastyle
        charstyle.priority = 100
        charstyle.hidden = False
        charstyle.locked = False
        charstyle.font.name = font
        charstyle.font.size = fsize
        charstyle.font.bold = bold
        charstyle.font.italic = italic
        charstyle.font.color.rgb = RGBColor(*rgb)
        charstyle.font.highlight_color = hilite
        return charstyle

    def _format_notation(self):
        #     section: Section = self.doc.sections[0]
        #     section.top_margin = Cm(1)
        #     section.bottom_margin = Cm(1)
        #     section.left_margin = Cm(0.5)
        #     section.right_margin = Cm(0.5)

        # Paragraph styles

        # self.basicparastyle = self._create_para_style(
        #     "BasicParagraph",
        #     keep_with_next=True,
        #     hanging_indent=Cm(2),
        #     tab_stops=[
        #         (Cm(11), WD_TAB_ALIGNMENT.RIGHT),
        #         (Cm(19), WD_TAB_ALIGNMENT.RIGHT),
        #     ],
        # )
        # self.tabbedparastyle = self._create_para_style(
        #     "TabbedParagraph",
        #     keep_with_next=True,
        #     tab_stops=[(Cm(2) + i * Cm(1.1), WD_TAB_ALIGNMENT.LEFT) for i in range(17)],
        # )
        # self.separatorparastyle = self._create_para_style(
        #     "SeparatorParagraph",
        # )

        self.basicparastyle: ParagraphStyle = self.doc.styles["GamelanBasicParagraph"]
        self.adjustrightparastyle: ParagraphStyle = self.doc.styles["GamelanAdjustRightParagraph"]
        self.separatorparastyle: ParagraphStyle = self.doc.styles["GamelanSeparatorParagraph"]
        self.tablestyle: ParagraphStyle = self.doc.styles["GamelanNotationTable"]
        self.basicparastyle.paragraph_format.tab_stops.add_tab_stop(self.TAG_COLWIDTH)

        # Character styles

        self.tagstyle = self._create_char_style("GamelanTag", base=self.tablestyle, fsize=Pt(8), bold=True)
        self.commentstyle = self._create_char_style("GamelanComment", fsize=Pt(9), rgb=(128, 128, 64))
        self.sequencestyle = self._create_char_style("GamelanSequence", fsize=Pt(9), rgb=(128, 52, 13))
        self.partstyle = self._create_char_style("GamelanPart", fsize=Pt(11), bold=True)
        self.metadatastyle = self._create_char_style("GamelanMetadata", fsize=Pt(8), italic=True)
        self.metadatastyle_hl = self._create_char_style("GamelanMetadataHilite", fsize=Pt(9), rgb=(0, 0, 255))
        self.labelstyle = self._create_char_style(
            "GamelanLabel", font="Courier New", base=self.tablestyle, fsize=Pt(9), bold=True, rgb=(0, 0, 255)
        )
        self.notationstyle = self._create_char_style(
            "GamelanNotation", base=self.tablestyle, font="Bali Music 5", fsize=Pt(9)
        )

    def textwidth(self, text: str, font: str, fontsize: int) -> Length:
        # TODO: copy all the font files in the code.
        # Will also require to instruct docx to use these files.
        """Determine the width of a text for the given font and size.
        Args:
            text (str): Text to evaluate.
            font (str): Font name.
            fontsize (int): Font size in pt.
        Raises:
            FileNotFoundError: If the font file was not found.

        Returns:
            Length:
        """
        try:
            fontfile = os.path.join(self.pdf_settings.folder, self.pdf_settings.fonts[font])
            imagefont = ImageFont.truetype(fontfile, fontsize)
        except:
            raise FileNotFoundError(f"Could not determine length of text={text} font={font} fontsize={fontsize}.")

        # Create a dummy image and draw object to determine text width
        img = Image.new("RGB", (1, 1))
        draw = ImageDraw.Draw(img)
        width_px = draw.textlength(text, font=imagefont, font_size=fontsize)
        width = Inches(width_px / 72)
        return width

    def _edit_header(self):
        last_modif_epoch = os.path.getmtime(self.run_settings.notation.notation_filepath)
        last_modif_date = datetime.fromtimestamp(last_modif_epoch).strftime("%d-%b-%Y")
        header_paragraph = self.doc.sections[0].header.paragraphs[0]
        title_run = next((run for run in header_paragraph.runs if run.text == "{title}"), None)
        date_run = next((run for run in header_paragraph.runs if run.text == "{date}"), None)
        if title_run:
            title_run.text = self.score.title
        if date_run:
            date_run.text = last_modif_date

    def _bali_music_5_width(self, notes: list[Note]) -> int:
        cellborders = Cm(0.2)
        if isinstance(notes, list):
            nrchars = len([note for note in notes if not note.stroke is Stroke.GRACE_NOTE])
        else:
            nrchars = notes
        width_to_height_ratio = 8 / 12
        font_height = Pt(9)
        width = int(nrchars * font_height * width_to_height_ratio + cellborders)
        compare_width = self.textwidth("a" * nrchars, "Bali Music 5", 9) + cellborders
        return width

    def _format_metadata_cells(self, row: _Row) -> None:
        for cell in row.cells:
            tcPr = cell._element.get_or_add_tcPr()
            noborder = """<w:tcBorders xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
                        <w:left w:val="nil"/>
                        <w:right w:val="nil"/>
                    </w:tcBorders>
                  """
            aligntop = '<w:vAlign w:val="top" xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"/>'
            tcPr.insert(0, parse_xml(noborder))
            tcPr.insert(0, parse_xml(aligntop))

    def _is_empty(self, row: _Row, range: list[int] | None) -> bool:
        """Determines if the row has content in the given range.
        Args:
            row (_Row): The row to check.
            first (int): Index of the first cell of the range.
            last (int): Index of the last cell in the range.
        Returns:
            bool: True if all cells in the range are empty, False otherwise.
        """
        for cell in row.cells[range[0] : range[1] + 1 if range[1] > 0 else range[1]]:
            if cell.text.strip():
                return False
        return True

    def _to_aggregated_tags(self, positions: list[Position]) -> list[str]:
        tags = set(pos.instrumenttype.lower() for pos in positions)
        gangsa = {InstrumentType.PEMADE.lower(), InstrumentType.KANTILAN.lower()}
        if gangsa.issubset(tags):
            tags = tags.difference(gangsa).union({"gangsa"})
        return tags

    def _new_metadata_row(self, table: Table, merged: list[int] = [], parastyle: ParagraphStyle = None) -> _Row:
        row = table.add_row()
        if merged:
            merged = row.cells[merged[0]].merge(row.cells[merged[1]])
        if parastyle:
            for cell in row.cells:
                cell.paragraphs[0].style = parastyle
        self._format_metadata_cells(row)
        return row

    def _default_formatter(
        self,
        value: typing.Any | None,
        meta: MetaDataBaseModel,
        before: str = "",
        after: str = "",
        paragraph: Paragraph = None,
        charstyle: CharacterStyle = None,
    ) -> None:
        """Simple formatter for metadata. Adds the value to the given paragraph.
        Args:
            value (typing.Any): Value that should be written. If missing, the metadata DEFAULTPARAM attribute will be used.
            meta (MetaDataBaseModel): Metadata object.
            before (str, optional): Text that should precede the value. Defaults to "".
            after (str, optional): Text that should follow the value. Defaults to "".
            paragraph (Paragraph, optional): The paragraph to which the value should be written. Defaults to None.
            charstyle (CharacterStyle, optional): Character style for the added text. Defaults to `metadatastyle`.
        """
        charstyle = charstyle or self.defaultcharstyle
        value_ = value or getattr(meta, meta.DEFAULTPARAM)
        paragraph.add_run(f"{before}{value_}{after}", style=charstyle)

    def _list_formatter(
        self,
        value: list[str],
        meta: MetaDataBaseModel,
        before: str = "",
        after: str = "",
        paragraph: Paragraph = None,
        charstyle: CharacterStyle = None,
    ) -> None:
        """Formatter for a metadata list value. Formats the value to a comma separated list.
            The list items are formatted using the `labelstyle` character style.
        Args:
            value (typing.Any): Value that should be written. If missing, the metadata DEFAULTPARAM attribute will be used.
            meta (MetaDataBaseModel): Metadata object.
            before (str, optional): Text that should precede the value. Defaults to "".
            after (str, optional): Text that should follow the value. Defaults to "".
            paragraph (Paragraph, optional): The paragraph to which the value should be written. Defaults to None.
            charstyle (CharacterStyle, optional): Character style for the added text. Defaults to `metadatastyle`.
        """
        charstyle = charstyle or self.defaultcharstyle
        value_ = value or getattr(meta, meta.DEFAULTPARAM)
        if not isinstance(value_, list):
            value_ = [value_]
        if before:
            paragraph.add_run(before, style=charstyle)
        paragraph.add_run(", ".join(value_), style=self.labelstyle)
        if after:
            paragraph.add_run(after, style=charstyle)

    def _gradual_change_formatter(
        self,
        value: list[str],
        meta: MetaDataBaseModel,
        before: str = "",
        after: str = "",
        paragraph: Paragraph = None,
        charstyle: CharacterStyle = None,
    ):
        """Formatter for GradualChangeMetadata subclasses. These values can include a range of beats to which the metadata
           applies. Called for TEMPO and DYNAMICS metadata which can gradually change over several beats.
           The value will be followed by a dashed line that will reach to the right margine of the (merged) cell.
        Args:
            value (typing.Any): Value that should be written. If missing, the metadata DEFAULTPARAM attribute will be used.
            meta (MetaDataBaseModel): Metadata object.
            before (str, optional): Text that should precede the value. Defaults to "".
            after (str, optional): Text that should follow the value. Defaults to "".
            paragraph (Paragraph, optional): The paragraph to which the value should be written. Defaults to None.
            charstyle (CharacterStyle, optional): Character style for the added text. Defaults to `metadatastyle`.
        """
        # If the metadata spans several beats, a tab character will be added before or after the value and a tab stop
        # with preceding dashes will be added at the right-hand margin of the merged cell.
        charstyle = charstyle or self.defaultcharstyle
        tab = "\t" if meta.beat_count > 1 else ""
        if meta.metatype == "TEMPO":
            if self.current_tempo == -1:
                self.current_tempo = meta.value
                return
            value = f"faster{tab}" if meta.value > self.current_tempo else f"slower{tab}"
            self.current_tempo = meta.value
        elif meta.metatype == "DYNAMICS":
            position_tags = self._to_aggregated_tags(meta.positions)
            value = f"{tab}{", ".join(position_tags)}{": " if position_tags else ""}{meta.abbreviation}"
            self.current_dynamics = meta.value
        cellwidth = paragraph._parent.width
        # Add a tab with preceding dashes to the end of the merged cells (need to create style for this).
        # The style name should be unique but paragraph.__hash__() seems not always to be unique.
        # Hence the additional random value.
        parastyle = self._create_para_style(
            f"paraformat_{paragraph.__hash__()}{randint(1, 1000000):07d}",
            tab_stops=[[cellwidth, WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.DASHES]],
        )
        # paragraph.paragraph_format.tab_stops.clear_all()
        paragraph.style = parastyle
        paragraph.add_run(value, charstyle)

    def _add_metadata(
        self,
        table: Table,
        metalist: list[MetaDataBaseModel],
        value: str = None,
        cellnr: int | str = 1,
        merge: int | str | None = None,
        mergetype: MergeType = MergeType.LAST_CELL,
        parastyle: ParagraphStyle = None,
        charstyle: CharacterStyle = None,
        before: str = "",
        after: str = "",
        formatter: typing.Callable = None,
    ):
        """Generic method that appends a list of metadata items of the same type to the gongan table.
           Note that table contains an overflow column (after the last beat column) to accommodate long metadata texts.
        Args:
            table (Table): The table to which to add the metadata.
            metalist (list[MetaData]): List of similar metadata items.
            value (str, optional): The value to print. Defaults to the default attrribute of the metadata item (set in DEFAULTPARAM).
            cellnr: (int, optional): The table cell in which to write the value. Defaults to 1.
            merged (list[int], optional): If true, all cells from cellnr to the end of the row will be merged. Defaults to True.
            parastyle (ParagraphStyle, optional): Paragraph style. Defaults to basicparastyle.
            charstyle (CharacterStyle, optional): Character style. Defaults to metadatastyle.
            before (str, optional): Text to print before the value. Defaults to "".
            after (str, optional): Text to print after the value. Defaults to "".
            formatter (Callable, optional): Function that formats and writes the text. Defaults to default_formatter.
        """
        for meta in metalist:
            parastyle = parastyle or self.defaultparastyle
            charstyle = charstyle or self.defaultcharstyle
            # value = getattr(meta, attr)
            # value = eval(attr) if isinstance(attr, str) else attr
            cellnr_ = getattr(meta, cellnr) if isinstance(cellnr, str) else cellnr
            if merge:
                merge_ = merge
                if isinstance(merge, str):
                    merge_ = getattr(meta, merge)
                    # table has len(gongan.beats) + 2. The column of the last beat is table.columns -2.
                    merge_ = len(table.columns) - 2 if merge_ < 0 else merge_
                if mergetype is MergeType.RANGE:
                    merge_range = None if merge_ <= 1 else [cellnr_, cellnr_ + merge_ - 1]
                elif mergetype is MergeType.LAST_CELL:
                    merge_range = [cellnr_, cellnr_ + merge_ - 1]
            else:
                merge_range = None
            # Don't add a new row if the content fits on the last table row,
            # except if all the row cells are merged (=intentional blank row).
            if (
                table.rows
                and not table.rows[-1].cells[0].grid_span == len(table.columns)
                and self._is_empty(table.rows[-1], merge_range or [cellnr_, cellnr_])
            ):
                row = table.rows[-1]
                if merge_range:
                    row.cells[merge_range[0]].merge(row.cells[merge_range[1]])
            else:
                row = self._new_metadata_row(table, merged=merge_range, parastyle=parastyle)
            paragraph = row.cells[cellnr_].paragraphs[0]
            formatter = formatter or self._default_formatter
            formatter(value, meta, before=before, after=after, paragraph=paragraph, charstyle=charstyle)
            if "passes" in meta.model_fields:
                if meta.passes:
                    if any(p for p in meta.passes if p > 0):
                        paragraph.add_run(f" (pass {','.join(str(p) for p in meta.passes)})", charstyle)
            # merge additional cells if necessary to accommondate the text width
            text_width = sum(
                self.textwidth(
                    run.text,
                    run.style.font.name or paragraph.style.font.name,
                    run.style.font.size.pt or paragraph.style.font.size.pt,
                )
                for run in paragraph.runs
            )
            delta = Cm(0.2)
            if (cwidth := paragraph._parent.width) < text_width + delta:
                # if text is in the last gongan column: merge the cell with the overflow cell to the right
                # Merge the cell with empty adjacent cells if possible until the required width is reached.
                if paragraph.style.paragraph_format.alignment == WD_ALIGN_PARAGRAPH.RIGHT:
                    adjacent_cell = cellnr_
                    while (
                        row.cells[cellnr_].width < text_width + delta
                        and adjacent_cell - 1 >= 0
                        and not row.cells[adjacent_cell - 1].text
                    ):
                        adjacent_cell -= 1
                        # TODO: if the adjcell is the rightmost cell of the table, adjust the column width if needed
                        row.cells[adjacent_cell].merge(row.cells[cellnr_])
                        cellnr_ = adjacent_cell
                else:
                    adjacent_cell = cellnr_ + row.cells[cellnr_].grid_span - 1
                    while (
                        row.cells[cellnr_].width < text_width + delta
                        and adjacent_cell + 1 <= table._column_count - 1
                        and not row.cells[adjacent_cell + 1].text
                    ):
                        adjacent_cell += 1
                        # TODO: if the adjcell is the rightmost cell of the table, adjust the column width if needed
                        row.cells[cellnr_].merge(row.cells[adjacent_cell])

    def _add_comments(self, table: Table, comments: list[str]):
        for comment in comments:
            if not comment.startswith("#"):
                row = self._new_metadata_row(table, merged=[1, -1], parastyle=self.basicparastyle)
                paragraph = row.cells[-1].paragraphs[0]
                paragraph.add_run(f"{comment}", self.commentstyle)

    def _add_metadata_before_gongan(self, table: Table, gongan: Gongan):
        """Adds the gongan metadata and comments that should appear at the top of the gongan.
        Only processes the metadata that is meaningful for the PDF notation. E.g. ValidationMeta,
        GonganMeta, AutoKempyungMeta, OctavateMeta are skipped.
        Args:
            table (Table): pre-formatted docx table with beats + 1 columns
            gongan (Gongan): the gongan to which the metadata belongs
        """
        metatypes = {meta.data.metatype for meta in gongan.metadata}
        metadict = {
            metatype: [meta.data for meta in gongan.metadata if meta.data.metatype == metatype]
            for metatype in metatypes
        }
        # still to add DynamicsMeta, KempliMeta, SuppressMeta, TempoMeta

        if metalist := metadict.get("PART", None):
            self._add_metadata(table, metalist, merge=-1, charstyle=self.partstyle)
        self._add_comments(table, gongan.comments)
        if metalist := metadict.get("LABEL", None):
            self._add_metadata(table, metalist, cellnr="beat", merge=-1, charstyle=self.labelstyle)

        if metalist := metadict.get("TEMPO", None):
            self._add_metadata(
                table,
                metalist,
                cellnr="first_beat",
                merge="beat_count",
                mergetype=MergeType.RANGE,
                parastyle=None,
                charstyle=self.metadatastyle,
                formatter=self._gradual_change_formatter,
            )

        if metalist := metadict.get("DYNAMICS", None):
            self._add_metadata(
                table,
                metalist,
                cellnr="first_beat",
                merge="beat_count",
                mergetype=MergeType.RANGE,
                parastyle=None,
                charstyle=self.metadatastyle,
                formatter=self._gradual_change_formatter,
            )

    def _add_metadata_after_gongan(self, table: Table, gongan: Gongan):
        """Adds the gongan metadata that should appear at the bottom of the gongan.
        Args:
            table (Table): pre-formatted docx table with beats + 1 columns
            gongan (Gongan): the gongan to which the metadata belongs
        """
        metatypes = {meta.data.metatype for meta in gongan.metadata}
        metadict = {
            metatype: [meta.data for meta in gongan.metadata if meta.data.metatype == metatype]
            for metatype in metatypes
        }
        # not added (unnecessary): WaitMeta
        if metalist := metadict.get("REPEAT", None):
            self._add_metadata(
                table,
                metalist,
                merge=len(gongan.beats),  # right-aligned: do not use extra column
                before="repeat ",
                after="X",
                parastyle=self.adjustrightparastyle,
                charstyle=self.metadatastyle_hl,
            )
        if metalist := metadict.get("GOTO", None):
            self._add_metadata(
                table,
                metalist,
                merge="from_beat",
                before="go to ",
                parastyle=self.adjustrightparastyle,
                charstyle=self.metadatastyle_hl,
                formatter=self._list_formatter,
            )
        if metalist := metadict.get("SEQUENCE", None):
            self._new_metadata_row(table, merged=[0, -1])
            self._add_metadata(
                table,
                metalist,
                merge=-1,
                before="sequence: ",
                parastyle=self.basicparastyle,
                charstyle=self.sequencestyle,
                formatter=self._list_formatter,
            )

    def _measure_to_str(self, notes: list[Note]) -> str:
        if not notes:
            return ""
        notechars = "".join(note.symbol for note in notes)
        return notechars

    def _clean_staves(self, gongan: Gongan) -> dict[Position, list[list[Note]]]:
        if not gongan.beats:
            return dict()
        staves = {
            position: [
                [note for note in beat.measures[position].passes[DEFAULT].notes if not note.autogenerated]
                for beat in gongan.beats
                # skip beats that are completely autogenerated (e.g. generated from WAIT metadata)
                if not all(
                    note.autogenerated for measure in beat.measures.values() for note in measure.passes[DEFAULT].notes
                )
            ]
            for position in gongan.beats[0].measures.keys()
        }
        # Remove empty staves (which contained only autogenerated notes)
        # staves = {position: [measure for measure in stave if measure] for position, stave in staves.items()}
        return staves

    def _add_staves(self, table: Table, staves: list[list[Note]], pos_tags: dict[Position, str]) -> None:
        for position in sorted(list(pos_tags.keys()), key=lambda x: x.sequence):
            cells: list[_Cell] = table.add_row().cells
            paragraph = cells[0].paragraphs[0]
            paragraph.add_run(pos_tags[position], self.tagstyle)
            for i, measure in enumerate(staves[position], 1):
                paragraph = cells[i].paragraphs[0]
                paragraph.add_run(self._measure_to_str(measure), self.notationstyle)

    def _create_table(self, gongan: Gongan) -> Table:
        """Creates a table for the gongan. Column 1 will contain the position tags. The last column is an overflow
           column for long metadata text values. Its width is maximized so that it extends to to the right margin.
        Args:
            gongan (Gongan): The gongan for which the table should be created.

        Returns:
            Table:
        """
        staves = self._clean_staves(gongan)
        # The number of beats in the cleaned staves might be less than len(gongan.beats)
        beat_count = len(list(staves.values())[0])
        beat_colwidths = [
            max(self._bali_music_5_width(staves[position][i]) for position in staves.keys()) for i in range(beat_count)
        ]
        colwidths = [self.TAG_COLWIDTH] + beat_colwidths + [Cm(0.1)]  # last column width will be set below.
        table = self.doc.add_table(rows=0, cols=len(colwidths), style=self.tablestyle)
        table.autofit = False
        for i in range(len(colwidths)):
            table.columns[i].width = colwidths[i]
        total_width = sum(colwidths)
        section = self.doc.sections[0]
        table.columns[i].width = max(0, section.page_width - section.left_margin - section.right_margin - total_width)
        return table

    def _convert_to_pdf(self) -> None:  # score_validation
        """Converts a score object to readable notation and saves it to PDF format (via DOCX).
        Args:
            score (Score): The score
        """
        METADATA_KEEP = ["DYNAMICS", "GOTO", "LABEL", "PART", "REPEAT", "SEQUENCE", "SUPPRESS", "TEMPO"]
        # score_dict = sum((gongan_to_records(gongan, score.settings.font.fontversion) for gongan in score.gongans), [])
        self._edit_header()
        # self.doc.add_heading(self.score.title)
        # self.doc.add_paragraph(style=self.separatorparastyle)
        for gongan in self.score.gongans:
            staves = self._clean_staves(gongan)
            table = self._create_table(gongan)
            self._add_metadata_before_gongan(table, gongan)
            if not gongan.beats:
                continue
            pos_tags = aggregate_positions(gongan)
            self._add_staves(table, staves, pos_tags)
            self._add_metadata_after_gongan(table, gongan)
            self.doc.add_paragraph(style=self.separatorparastyle)

    def is_docfile_closed():
        return

    @ParserModel.main
    def create_notation(self):
        if self.pickle:
            with open(self.filepath.replace("docx", "pickle"), "wb") as picklefile:
                pickle.dump(self.score, picklefile)

        self._convert_to_pdf()
        self.doc.save(self.filepath)
        docx2pdf.convert(self.filepath, self.filepath.replace("docx", "pdf"))
        self.logger.info(f"Notation file saved as {self.filepath}")


if __name__ == "__main__":
    # print(ScoreToPDFConverter._bali_music_5_width(40) / Cm(1))
    # print(textwidth("This is a very long sentence, try this!", font="Arial", fontsize=10).cm)
    # exit()

    class Source(Enum):
        SINOMLADRANG = {
            "folder": r"C:\Users\marcp\Documents\administratie\_VRIJETIJD_REIZEN\Scripts-Programmas\PythonProjects\gamelan-notation\data\notation\sinom ladrang",
            "filename": "Sinom Ladrang_full_GAMELAN1.pickle",
        }
        LEGONGMAHAWIDYA = {
            "folder": r"C:\Users\marcp\Documents\administratie\_VRIJETIJD_REIZEN\Scripts-Programmas\PythonProjects\gamelan-notation\data\notation\legong mahawidya",
            "filename": "Legong Mahawidya_full_GAMELAN1.pickle",
        }
        SEKARGENDOT = {
            "folder": r"C:\Users\marcp\Documents\administratie\_VRIJETIJD_REIZEN\Scripts-Programmas\PythonProjects\gamelan-notation\data\notation\sekar gendot",
            "filename": "Sekar Gendot_full_GAMELAN1.pickle",
        }
        LENGKER = {
            "folder": r"C:\Users\marcp\Documents\administratie\_VRIJETIJD_REIZEN\Scripts-Programmas\PythonProjects\gamelan-notation\data\notation\lengker",
            "filename": "Lengker_full_GAMELAN1.pickle",
        }

    source = Source.LEGONGMAHAWIDYA

    path = os.path.join(source.value["folder"], source.value["filename"])
    with open(path, "rb") as picklefile:
        score = pickle.load(picklefile)
        ScoreToPDFConverter(score, False).create_notation()
