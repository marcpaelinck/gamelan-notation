import os
import pickle
from os import path

import docx2pdf
from docx import Document
from docx.document import Document as DocxDocument
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_COLOR_INDEX
from docx.oxml.parser import parse_xml
from docx.oxml.xmlchemy import BaseOxmlElement
from docx.section import Section
from docx.shared import Cm, Pt, RGBColor
from docx.styles.style import BaseStyle, CharacterStyle, ParagraphStyle, _TableStyle
from docx.table import Table, _Cell, _Row
from docx.text.paragraph import Paragraph

from src.common.classes import Gongan, Note, Score
from src.common.constants import DEFAULT, Position, Stroke
from src.common.metadata_classes import MetaData
from src.notation2midi.score_to_notation import aggregate_positions

# The following functions save a human readable version of the score to a PDF file.


class ScoreToPDFConverter:
    TAG_COLWIDTH = Cm(2.3)

    def __init__(self, score: Score, pickle: bool = True):
        self.doc: DocxDocument = Document("./data/doc_template/Template.docx")
        self.score = score
        self.pickle = pickle
        self._format_notation()
        fpath, ext = os.path.splitext(score.settings.notation.midi_out_filepath)
        self.filepath = fpath + ".docx"
        self.current_para = None

    def _create_para_style(
        self,
        name: str,
        space_after: int = 0,
        keep_with_next: bool = False,
        hanging_indent: int = None,
        tab_stops: list[tuple[int, WD_STYLE_TYPE]] = [],
        align: WD_ALIGN_PARAGRAPH = WD_ALIGN_PARAGRAPH.LEFT,
    ) -> ParagraphStyle:
        parastyle: ParagraphStyle = self.doc.styles.add_style(name, style_type=WD_STYLE_TYPE.PARAGRAPH, builtin=False)
        parastyle.base_style = self.doc.styles["Normal"]
        parastyle.font.name = "Arial"
        parastyle.paragraph_format.space_after = space_after
        parastyle.paragraph_format.keep_with_next = keep_with_next
        parastyle.paragraph_format.tab_stops.clear_all()
        if hanging_indent:
            parastyle.paragraph_format.left_indent = hanging_indent
            parastyle.paragraph_format.first_line_indent = -hanging_indent
        for tab, tabtype in tab_stops:
            parastyle.paragraph_format.tab_stops.add_tab_stop(tab, alignment=tabtype)
        parastyle.font.highlight_color = WD_COLOR_INDEX.WHITE
        parastyle.paragraph_format.alignment = align

        return parastyle

    def _create_char_style(
        self,
        name: str,
        base: ParagraphStyle = None,
        font: str = "Arial",
        fsize: int = 10,
        bold=False,
        italic=False,
        rgb: tuple[int] = (0, 0, 0),
        hilite: WD_COLOR_INDEX = WD_COLOR_INDEX.WHITE,
    ) -> CharacterStyle:
        charstyle: CharacterStyle = self.doc.styles.add_style(name, style_type=WD_STYLE_TYPE.CHARACTER, builtin=False)
        charstyle.base_style = base or self.basicparastyle
        charstyle.priority = 100
        charstyle.hidden = False
        charstyle.locked = False
        charstyle.font.name = font
        charstyle.font.size = fsize
        charstyle.font.bold = bold
        charstyle.font.italic = italic
        charstyle.font.color.rgb = RGBColor(*rgb)
        charstyle.font.highlight_color = hilite
        return charstyle

    def _format_notation(self):
        section: Section = self.doc.sections[0]
        section.top_margin = Cm(1)
        section.bottom_margin = Cm(1)
        section.left_margin = Cm(0.5)
        section.right_margin = Cm(0.5)

        # Paragraph styles

        # self.basicparastyle = self._create_para_style(
        #     "BasicParagraph",
        #     keep_with_next=True,
        #     hanging_indent=Cm(2),
        #     tab_stops=[
        #         (Cm(11), WD_TAB_ALIGNMENT.RIGHT),
        #         (Cm(19), WD_TAB_ALIGNMENT.RIGHT),
        #     ],
        # )
        # self.tabbedparastyle = self._create_para_style(
        #     "TabbedParagraph",
        #     keep_with_next=True,
        #     tab_stops=[(Cm(2) + i * Cm(1.1), WD_TAB_ALIGNMENT.LEFT) for i in range(17)],
        # )
        # self.separatorparastyle = self._create_para_style(
        #     "SeparatorParagraph",
        # )

        self.basicparastyle: ParagraphStyle = self.doc.styles["GamelanBasicParagraph"]
        self.adjustrightparastyle: ParagraphStyle = self.doc.styles["GamelanAdjustRightParagraph"]
        self.separatorparastyle: ParagraphStyle = self.doc.styles["GamelanSeparatorParagraph"]
        self.tablestyle: ParagraphStyle = self.doc.styles["GamelanNotationTable"]
        self.basicparastyle.paragraph_format.tab_stops.add_tab_stop(self.TAG_COLWIDTH)

        # Character styles

        self.tagstyle = self._create_char_style("GamelanTag", base=self.tablestyle, fsize=Pt(8), bold=True)
        self.commentstyle = self._create_char_style("GamelanComment", fsize=Pt(9), rgb=(128, 128, 64))
        self.partstyle = self._create_char_style("GamelanPart", fsize=Pt(11), bold=True)
        self.metadatastyle = self._create_char_style("GamelanMetadata", fsize=Pt(8), italic=True)
        self.metadatastyle_hl = self._create_char_style("GamelanMetadataHilite", fsize=Pt(9), rgb=(0, 0, 255))
        self.labelstyle = self._create_char_style(
            "GamelanLabel", font="Courier New", base=self.tablestyle, fsize=Pt(9), bold=True, rgb=(0, 0, 255)
        )
        self.notationstyle = self._create_char_style(
            "GamelanNotation", base=self.tablestyle, font="Bali Music 5", fsize=Pt(9)
        )

    @classmethod
    def _bali_music_5_width(self, notes: list[Note]) -> int:
        if isinstance(notes, list):
            nrchars = len([note for note in notes if not note.stroke is Stroke.GRACE_NOTE])
        else:
            nrchars = notes
        width_to_height_ratio = 8 / 12
        correction = 6 / 6
        font_height = Pt(9)
        return int(nrchars * font_height * width_to_height_ratio * correction + Cm(0.2))

    def _format_metadata_cells(self, row: _Row) -> None:
        for cell in row.cells:
            tcPr = cell._element.get_or_add_tcPr()
            noborder = """<w:tcBorders xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
                        <w:left w:val="nil"/>
                        <w:right w:val="nil"/>
                    </w:tcBorders>
                  """
            aligntop = '<w:vAlign w:val="top" xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"/>'
            tcPr.insert(0, parse_xml(noborder))
            tcPr.insert(0, parse_xml(aligntop))

    def _add_labels(
        self, labels: list[str], separator: str = ", ", before: str = "", after: str = "", paragraph: Paragraph = None
    ) -> None:
        if before:
            paragraph.add_run(before, self.metadatastyle_hl)
        for count, label in enumerate(labels, 1):
            paragraph.add_run(label, self.labelstyle)
            # Add a separator between the labels
            if count < len(labels):
                paragraph.add_run(separator, self.metadatastyle_hl)
        if after:
            paragraph.add_run(after, self.metadatastyle_hl)

    def _add_metadata_row(self, table: Table, merged: list[int] = [], style: BaseStyle = None) -> _Row:
        row = table.add_row()
        if merged:
            merged = row.cells[merged[0]].merge(row.cells[merged[1]])
        if style:
            for cell in row.cells:
                cell.paragraphs[0].style = style
        self._format_metadata_cells(row)
        return row

    def _add_comments(self, table: Table, comments: list[str]):
        for comment in comments:
            if not comment.startswith("#"):
                row = self._add_metadata_row(table, merged=[1, -1], style=self.basicparastyle)
                paragraph = row.cells[-1].paragraphs[0]
                paragraph.add_run(f"{comment}", self.commentstyle)

    def add_goto_meta(self, table: Table, metalist: list[MetaData]) -> None:
        for meta in metalist:
            row = self._add_metadata_row(table, merged=[1, -1], style=self.adjustrightparastyle)
            paragraph = row.cells[-1].paragraphs[0]
            self._add_labels([meta.label], before=f"go to ", paragraph=paragraph)
            if meta.passes:
                paragraph.add_run(
                    f" (pass {','.join(str(p) for p in meta.passes)})",
                    self.metadatastyle_hl,
                )

    def add_part_meta(self, table: Table, metalist: list[MetaData]) -> None:
        for meta in metalist:
            row = self._add_metadata_row(table, merged=[0, -1], style=self.basicparastyle)
            paragraph = row.cells[-1].paragraphs[0]
            paragraph.add_run(f"\t{meta.name}", self.partstyle)

    def add_label_meta(self, table: Table, metalist: list[MetaData]) -> None:
        for meta in metalist:
            row = self._add_metadata_row(table, merged=[meta.beat, -1], style=self.basicparastyle)
            paragraph = row.cells[meta.beat].paragraphs[0]
            paragraph.add_run(f"{meta.name}", self.labelstyle)

    def add_sequence_meta(self, table: Table, metalist: list[MetaData]) -> None:
        # skip a row
        self._add_metadata_row(table, style=self.basicparastyle)
        for meta in metalist:
            row = self._add_metadata_row(table, merged=[1, -1], style=self.basicparastyle)
            row.cells[0].paragraphs[0].add_run("sequence:", self.metadatastyle_hl)
            paragraph = row.cells[1].paragraphs[0]
            self._add_labels([str(p) for p in meta.value], paragraph=paragraph)

    def _add_metadata_gongan_before(self, table: Table, gongan: Gongan):
        # Add comments, except the ones starting with # (starting with ## in notation)
        metatypes = {meta.data.metatype for meta in gongan.metadata}
        metadict = {
            metatype: [meta.data for meta in gongan.metadata if meta.data.metatype == metatype]
            for metatype in metatypes
        }
        # not added (unnecessary): ValidationMeta, GonganMeta, AutoKempyungMeta, OctavateMeta
        # still to add DynamicsMeta, KempliMeta, SuppressMeta, TempoMeta

        if metalist := metadict.get("PART", None):
            self.add_part_meta(table, metalist)
        self._add_comments(table, gongan.comments)
        if metalist := metadict.get("LABEL", None):
            self.add_label_meta(table, metalist)
        # if metalist := metadict.get("TEMPO", None):
        #     for meta in metalist:
        #         term = "faster" if meta.value > gongan.beats[meta.first_beat - 1].bpm_start[DEFAULT] else "slower"
        #         paragraph = self.doc.add_paragraph(style=self.basicparastyle)
        #         paragraph.add_run("\t" * (meta.first_beat - 1) + term, self.metadatastyle)

    def _add_metadata_gongan_after(self, table: Table, gongan: Gongan):
        metatypes = {meta.data.metatype for meta in gongan.metadata}
        metadict = {
            metatype: [meta.data for meta in gongan.metadata if meta.data.metatype == metatype]
            for metatype in metatypes
        }
        # not added (unnecessary): WaitMeta
        tolastbeat = "\t" * (3 if len(gongan.beats) > 8 or (len(gongan.beats) > 4 and gongan.beat_duration > 4) else 2)

        if metalist := metadict.get("GOTO", None):
            self.add_goto_meta(table, metalist)

        if metalist := metadict.get("REPEAT", None):
            for meta in metalist:
                paragraph = self.doc.add_paragraph(style=self.basicparastyle)
                paragraph.add_run(f"{tolastbeat}repeat {meta.count}", self.metadatastyle_hl)
        if metalist := metadict.get("SEQUENCE", None):
            self.add_sequence_meta(table, metalist)

    def _measure_to_str(self, notes: list[Note]) -> str:
        if not notes:
            return ""
        notechars = "".join(note.symbol for note in notes)
        return notechars

    def _clean_staves(self, gongan: Gongan) -> dict[Position, list[list[Note]]]:
        if not gongan.beats:
            return dict()
        staves = {
            position: [
                [note for note in beat.measures[position].passes[DEFAULT].notes if not note.autogenerated]
                for beat in gongan.beats
                # skip beats that are completely autogenerated (e.g. generated from WAIT metadata)
                if not all(
                    note.autogenerated for measure in beat.measures.values() for note in measure.passes[DEFAULT].notes
                )
            ]
            for position in gongan.beats[0].measures.keys()
        }
        # Remove empty staves (which contained only autogenerated notes)
        # staves = {position: [measure for measure in stave if measure] for position, stave in staves.items()}
        return staves

    def _add_staves(self, table: Table, staves: list[list[Note]], pos_tags: dict[Position, str]) -> None:
        for position in sorted(list(pos_tags.keys()), key=lambda x: x.sequence):
            cells: list[_Cell] = table.add_row().cells
            paragraph = cells[0].paragraphs[0]
            paragraph.add_run(pos_tags[position], self.tagstyle)
            for i, measure in enumerate(staves[position], 1):
                paragraph = cells[i].paragraphs[0]
                paragraph.add_run(self._measure_to_str(measure), self.notationstyle)

    def _create_table(self, gongan: Gongan) -> Table:
        staves = self._clean_staves(gongan)
        # The number of beats in the cleaned staves might be less than len(gongan.beats)
        beat_count = len(list(staves.values())[0])
        colwidths = [
            max(self._bali_music_5_width(staves[position][i]) for position in staves.keys()) for i in range(beat_count)
        ]
        colwidths = [self.TAG_COLWIDTH] + colwidths
        table = self.doc.add_table(rows=0, cols=beat_count + 1, style=self.tablestyle)
        table.autofit = False
        for i in range(len(colwidths)):
            table.columns[i].width = colwidths[i]
        print(f"gongan[{gongan.id}]: {[round(col.width/Cm(1), 2) for col in table.columns]}")
        return table

    def _convert_to_pdf(self) -> None:  # score_validation
        """Converts a score object to readable notation and saves it to PDF format (via DOCX).
        Args:
            score (Score): The score
        """
        METADATA_KEEP = ["DYNAMICS", "GOTO", "LABEL", "PART", "REPEAT", "SEQUENCE", "SUPPRESS", "TEMPO"]
        # score_dict = sum((gongan_to_records(gongan, score.settings.font.fontversion) for gongan in score.gongans), [])
        self.doc.add_heading(self.score.title)
        self.doc.add_paragraph(style=self.separatorparastyle)
        for gongan in self.score.gongans:
            staves = self._clean_staves(gongan)
            table = self._create_table(gongan)
            self._add_metadata_gongan_before(table, gongan)
            if not gongan.beats:
                continue
            pos_tags = aggregate_positions(gongan)
            self._add_staves(table, staves, pos_tags)
            self._add_metadata_gongan_after(table, gongan)
            self.doc.add_paragraph(style=self.separatorparastyle)

    def is_docfile_closed():
        return

    def save(self):
        self._convert_to_pdf()
        self.doc.save(self.filepath)
        # docx2pdf.convert(self.filepath)
        if self.pickle:
            with open(self.filepath.replace("docx", "pickle"), "wb") as picklefile:
                pickle.dump(self.score, picklefile)


if __name__ == "__main__":
    # print(ScoreToPDFConverter._bali_music_5_width(40) / Cm(1))
    folder = r"C:\Users\marcp\Documents\administratie\_VRIJETIJD_REIZEN\Scripts-Programmas\PythonProjects\gamelan-notation\data\notation\legong mahawidya"
    # filename = "Legong Mahawidya_full_GAMELAN1.pickle"
    filename = "Legong Mahawidya_full_GAMELAN1.pickle"
    path = os.path.join(folder, filename)
    with open(path, "rb") as picklefile:
        score = pickle.load(picklefile)
        ScoreToPDFConverter(score, False).save()
